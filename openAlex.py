import requests
import json
import os
import time
import datetime
import sys
import re
import urllib.parse
from dotenv import load_dotenv

# Import the docx libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu

# Load environment variables from .env file
load_dotenv()

# --- Utility Functions ---

def validate_api_key():
    """Validate that the Amplify API key is available"""
    API_KEY = os.getenv("AMPLIFY_API_KEY")
    if not API_KEY:
        print("Error: AMPLIFY_API_KEY not found in environment variables.")
        print("Please set your API key in a .env file or environment variable.")
        return None
    return API_KEY


def make_llm_request(messages, model, temperature, max_tokens):
    """
    Makes a chat request to the Amplify API using the correct payload structure.
    """
    url = "https://prod-api.vanderbilt.ai/chat"
    API_KEY = validate_api_key()
    if not API_KEY:
        return None
    
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"}

    prompt_text = next((msg["content"] for msg in messages if msg["role"] == "user"), "")
    
    payload = {
        "data": {
            "temperature": temperature,
            "max_tokens": max_tokens,
            "dataSources": [],
            "messages": messages,
            "options": {
                "ragOnly": False,
                "skipRag": True,
                "model": {"id": model},
                "prompt": prompt_text,
            },
        }
    }

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=60)
        
        if response.status_code == 200:
            return response.json()
        elif response.status_code == 401:
            print("Error: Unauthorized - Check your API key")
        elif response.status_code == 403:
            print("Error: Forbidden - API key may be invalid or expired")
        elif response.status_code >= 500:
            print(f"Error: Server error (HTTP {response.status_code}) - Please try again later")
        else:
            print(f"Error: Request failed with status code {response.status_code}")
            print(f"Response: {response.text}")
        return None
        
    except requests.exceptions.RequestException as e:
        print(f"Error: Request failed - {e}")
        return None

def openalex_search(query, limit=5, mailto=None):
    """
    Performs a search on the OpenAlex API for scholarly works.
    Returns a list of work objects.
    """
    base_url = "https://api.openalex.org/works"
    
    # Use mailto for polite pool and better performance
    params = {
        'search': query,
        'per-page': limit,
        'select': 'id,doi,title,authorships,cited_by_count,publication_year'
    }
    
    if mailto:
        params['mailto'] = mailto

    try:
        response = requests.get(base_url, params=params, timeout=15)
        response.raise_for_status()
        data = response.json()
        return data.get('results', [])
    except requests.exceptions.RequestException as e:
        print(f"Error: OpenAlex API request failed - {e}")
        return []

def generate_subcategories(topic):
    """
    Generate subcategories for a given research topic
    """
    print("Step 1: Generating subcategories...")
    system_message = "You are a helpful research assistant. Your task is to break down a broad topic into 3 to 5 key subcategories. Respond with a numbered list. Do not include any other text, explanations, or markdown fences."
    
    messages = [
        {"role": "user", "content": f"Given the research topic: '{topic}', generate a list of 3 to 5 logical subcategories for an in-depth research paper."}
    ]
    
    response = make_llm_request(
        messages=messages,
        model="gpt-4o-mini",
        temperature=0.4,
        max_tokens=200
    )
    
    if response:
        raw_response = response.get("data", "").strip()
        subcategories = []
        for line in raw_response.split('\n'):
            line = line.strip()
            match = re.match(r'^\d+\.\s*(.*?)(?::.*)?$', line)
            if match:
                subcategory = match.group(1).strip()
                subcategory = subcategory.strip('**')
                subcategories.append(subcategory)

        if subcategories:
            print(f"✅ Subcategories generated: {subcategories}")
            return subcategories
        else:
            print("❌ Failed to parse subcategories. The response was not a valid list.")
    return None

def research_subcategory(subcategory):
    """
    Researches a subcategory by using the OpenAlex API to find and summarize papers.
    It returns a tuple of (body_text, references_list).
    """
    print(f"\nStep 2: Researching subcategory: '{subcategory}' with OpenAlex...")
    
    # Use a specific email for the polite pool, if available
    email = os.getenv("OPENALEX_EMAIL")
    search_query = f"{subcategory} research paper"
    search_results = openalex_search(search_query, limit=5, mailto=email)
    
    if not search_results:
        print(f"❌ No scholarly articles found for '{subcategory}'.")
        return None, []
        
    context_data = ""
    references = []
    
    for result in search_results:
        title = result.get('title')
        
        authorships = result.get('authorships', [])
        author_names = ", ".join([a.get('author', {}).get('display_name') for a in authorships])
        doi = result.get('doi')
        
        # Format the reference string
        ref_string = f"{author_names} ({result.get('publication_year', 'n.d.')}). {title}."
        if doi:
            ref_string += f" {doi}"
        references.append(ref_string)

        context_data += f"Title: {title}\nAuthors: {author_names}\nDOI: {doi}\n\n"
    
    prompt = f"""Synthesize a detailed academic summary for the subtopic '{subcategory}' using the following scholarly article metadata.
    Trace connections between the articles and the broader subtopic.
    Your summary must be in plain text and should not contain any in-text citations or references.
    
    Scholarly Article Metadata:
    {context_data}
    """
    
    messages = [
        {"role": "user", "content": prompt}
    ]
    
    response = make_llm_request(
        messages=messages,
        model="gpt-4o",
        temperature=0.5,
        max_tokens=2500
    )
    
    if response:
        body_text = response.get("data", "").strip()
        print(f"✅ Research for '{subcategory}' completed.")
        return body_text, references
    
    print(f"❌ No research data was collected for '{subcategory}'.")
    return None, []

def generate_txt_document(main_topic, research_findings, all_references):
    """Generates and saves a plain text document."""
    print("\nStep 3: Coalescing research into a final TXT paper...")

    full_research_text = "\n\n".join(research_findings)

    prompt = f"""Write a full research paper on the topic: '{main_topic}'.
    Synthesize the following research findings for the content of the paper, tracing connections between the subtopics.
    Your output should be a single, plain text document without any special formatting like Markdown. Do not include a references section.
    
    Research Findings:
    {full_research_text}
    
    The paper should have the following sections:
    - Title
    - Abstract
    - Introduction
    - Body Sections (based on the research findings)
    - Conclusion
    """

    messages = [
        {"role": "user", "content": prompt}
    ]
    
    response = make_llm_request(messages=messages, model="gpt-4o", temperature=0.6, max_tokens=4000)
    
    if response:
        plain_text_paper = response.get("data", "")
        if not plain_text_paper:
            print("❌ The AI did not return any text for the paper.")
            return None
            
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"research_paper_{timestamp}.txt"
        
        unique_references = sorted(list(set(ref.strip() for ref in all_references if ref.strip())))
        final_text = plain_text_paper + "\n\nReferences\n\n" + "\n".join(unique_references)
        
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write(final_text)
        
        print(f"✅ Research paper generated and saved to '{output_filename}'")
        return output_filename
    else:
        print("❌ Failed to generate the final research paper.")
        return None
    
def generate_markdown_document(main_topic, research_findings, all_references):
    """Generates and saves a Markdown document."""
    print("\nStep 3: Coalescing research into a final Markdown paper...")

    full_research_text = "\n\n".join(research_findings)

    prompt = f"""Write a full research paper on the topic: '{main_topic}'.
    Synthesize the following research findings, using Markdown syntax to format the document. Do not include a references section.
    
    Use a single '#' for the main title, '##' for main sections (e.g., Introduction, Conclusion), and '###' for sub-sections.
    
    Research Findings:
    {full_research_text}
    """

    messages = [
        {"role": "user", "content": prompt}
    ]
    
    response = make_llm_request(messages=messages, model="gpt-4o", temperature=0.6, max_tokens=4000)
    
    if response:
        markdown_paper = response.get("data", "")
        if not markdown_paper:
            print("❌ The AI did not return any Markdown text for the paper.")
            return None
            
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"research_paper_{timestamp}.md"
        
        unique_references = sorted(list(set(ref.strip() for ref in all_references if ref.strip())))
        references_markdown = "\n\n## References\n\n" + "\n".join([f"- {ref}" for ref in unique_references])
        final_text = markdown_paper + references_markdown
        
        with open(output_filename, 'w', encoding='utf-8') as f:
            f.write(final_text)
        
        print(f"✅ Research paper generated and saved to '{output_filename}'")
        return output_filename
    else:
        print("❌ Failed to generate the final research paper.")
        return None
    

def generate_docx_document(main_topic, research_findings, all_references):
    """
    Coalesce research findings into a complete research paper and save as a .docx file.
    Applies APA formatting using python-docx.
    """
    print("\nStep 3: Coalescing research into a final DOCX paper...")
    
    full_research_text = "\n\n".join(research_findings)

    prompt = f"""Write a full research paper in APA format on the topic: '{main_topic}'.
    Synthesize the following research findings, tracing connections between the subtopics.
    The output should be a single, plain text document with no special formatting. Ensure section titles are on their own lines. Do not include a references section.
    
    Research Findings:
    {full_research_text}
    
    The paper should have the following sections, with each title on a new line:
    - Title
    - Abstract
    - Introduction
    - Body Sections (based on the research findings)
    - Conclusion
    """
    
    messages = [
        {"role": "user", "content": prompt}
    ]
    
    response = make_llm_request(messages=messages, model="gpt-4o", temperature=0.6, max_tokens=4000)
    
    if response:
        plain_text_paper = response.get("data", "")
        if not plain_text_paper:
            print("❌ The AI did not return any text for the paper.")
            return None
            
        doc = Document()
        
        # Set basic APA document formatting
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Function to find and add headings
        def add_dynamic_headings(doc, text):
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.lower() in ['title', 'abstract', 'introduction', 'conclusion']:
                    heading = doc.add_heading(line, level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif re.match(r'^[A-Z][a-zA-Z\s,]+$', line):
                    heading = doc.add_heading(line, level=2)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.line_spacing = 2.0
                    p.paragraph_format.first_line_indent = Inches(0.5)

        # Add body text sections
        add_dynamic_headings(doc, plain_text_paper)
        
        # Consolidate, deduplicate, and alphabetize the references
        unique_references = sorted(list(set(ref.strip() for ref in all_references if ref.strip())))
        
        # Add References section
        if unique_references:
            refs_heading = doc.add_heading("References", level=1)
            refs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for ref_entry in unique_references:
                p = doc.add_paragraph(ref_entry)
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Save the document
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"research_paper_{timestamp}.docx"
        doc.save(output_filename)
        
        print(f"✅ Research paper generated and saved to '{output_filename}'")
        return output_filename
    else:
        print("❌ Failed to generate the final research paper.")
        return None
    

def main():
    try:
        if not validate_api_key():
            sys.exit(1)

        main_topic = input("Enter the research topic: ")
        if not main_topic:
            print("No topic entered. Exiting.")
            sys.exit(1)

        doc_format = input("Enter the desired document format (txt, md, or docx): ").lower()
        if doc_format not in ['txt', 'md', 'docx']:
            print("Invalid format. Please choose 'txt', 'md', or 'docx'. Exiting.")
            sys.exit(1)
        
        # Step 1: Generate subcategories
        subcategories = generate_subcategories(main_topic)
        if not subcategories:
            print("❌ Cannot proceed without subcategories.")
            sys.exit(1)
        
        # Step 2: Research each subcategory and collect references
        all_research_findings = []
        all_references = []
        for subcategory in subcategories:
            research_text, references = research_subcategory(subcategory)
            if research_text:
                all_research_findings.append(research_text)
                all_references.extend(references)
        
        if not all_research_findings:
            print("❌ No research data was collected. Exiting.")
            sys.exit(1)

        # Step 3: Coalesce into a final document based on user input
        if doc_format == 'txt':
            final_doc_path = generate_txt_document(main_topic, all_research_findings, all_references)
        elif doc_format == 'md':
            final_doc_path = generate_markdown_document(main_topic, all_research_findings, all_references)
        elif doc_format == 'docx':
            final_doc_path = generate_docx_document(main_topic, all_research_findings, all_references)
        
        if final_doc_path:
            print(f"\n📄 Final research paper saved to '{final_doc_path}'")
        else:
            sys.exit(1)

    except KeyboardInterrupt:
        print("\nOperation cancelled by user.")
        sys.exit(0)
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()